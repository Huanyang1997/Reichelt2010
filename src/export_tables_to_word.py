#!/usr/bin/env python3
"""Export Table 5 and Table 8 regression results to a paper-style Word document."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ModelKey = Tuple[int, str]
MODEL_ORDER: List[ModelKey] = [
    (1, "model1"),
    (1, "model2"),
    (1, "model3"),
    (2, "model1"),
    (2, "model2"),
    (2, "model3"),
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Export regression tables to a Word document.")
    p.add_argument("--table5-fullcoef", default="outputs/table5_replication_fullcoef.csv")
    p.add_argument("--table8-fullcoef", default="outputs/table8_replication_fullcoef.csv")
    p.add_argument("--output-docx", default="outputs/table5_table8_replication_paperstyle.docx")
    return p.parse_args()


def fmt_est(x) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.3f}"


def fmt_p(x) -> str:
    if pd.isna(x):
        return ""
    x = float(x)
    if x < 0.001:
        return "<0.001"
    return f"{x:.3f}"


def model_stat_t5(df: pd.DataFrame, d: int, m: str) -> Tuple[float, float, float]:
    sub = df[(df["definition"] == d) & (df["model"] == m)]
    if sub.empty:
        return float("nan"), float("nan"), float("nan")
    one = sub.iloc[0]
    return one.get("f_value", float("nan")), one.get("f_pvalue", float("nan")), one.get("adj_r2", float("nan"))


def model_stat_t8(df: pd.DataFrame, d: int, m: str) -> Tuple[float, float, float]:
    sub = df[(df["definition"] == d) & (df["model"] == m)]
    if sub.empty:
        return float("nan"), float("nan"), float("nan")
    one = sub.iloc[0]
    return one.get("lr_stat", float("nan")), one.get("lr_pvalue", float("nan")), one.get("pseudo_r2", float("nan"))


def get_term(df: pd.DataFrame, d: int, m: str, term: str) -> Tuple[float, float]:
    sub = df[(df["definition"] == d) & (df["model"] == m) & (df["term"] == term)]
    if sub.empty:
        return float("nan"), float("nan")
    one = sub.iloc[0]
    return one.get("coef", float("nan")), one.get("pvalue", float("nan"))


def row_values(df: pd.DataFrame, term_def1: str | None, term_def2: str | None) -> List[str]:
    vals: List[str] = []
    for d, m in MODEL_ORDER:
        target = term_def1 if d == 1 else term_def2
        if target is None:
            vals.extend(["", ""])
            continue
        coef, pv = get_term(df, d, m, target)
        vals.extend([fmt_est(coef), fmt_p(pv)])
    return vals


def apply_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = None


def center_cell(cell) -> None:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_block(
    doc: Document,
    title: str,
    subtitle: str,
    rows: List[Tuple[str, List[str]]],
    note: str,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True

    n_rows = 3 + len(rows)
    n_cols = 13
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row 0: definition groups
    t.cell(0, 0).text = ""
    t.cell(0, 1).merge(t.cell(0, 6)).text = "Auditor Industry Specialist Definition 1"
    t.cell(0, 7).merge(t.cell(0, 12)).text = "Auditor Industry Specialist Definition 2"

    # Header row 1: model groups
    t.cell(1, 0).text = ""
    for idx, label in zip([1, 3, 5, 7, 9, 11], ["Model 1", "Model 2", "Model 3", "Model 1", "Model 2", "Model 3"]):
        t.cell(1, idx).merge(t.cell(1, idx + 1)).text = label

    # Header row 2: estimate / p-value
    t.cell(2, 0).text = "Variable"
    for c in [1, 3, 5, 7, 9, 11]:
        t.cell(2, c).text = "Estimate"
        t.cell(2, c + 1).text = "p-value"

    for r in range(3):
        for c in range(n_cols):
            center_cell(t.cell(r, c))

    for ridx, (label, vals) in enumerate(rows, start=3):
        t.cell(ridx, 0).text = label
        for j, v in enumerate(vals, start=1):
            t.cell(ridx, j).text = v
            center_cell(t.cell(ridx, j))

    p3 = doc.add_paragraph(note)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_table5_rows(df: pd.DataFrame) -> List[Tuple[str, List[str]]]:
    spec = [
        ("Intercept", "const", "const"),
        ("SIZE", "size", "size"),
        ("sigma(CFO)", "sigma_cfo", "sigma_cfo"),
        ("CFO", "cfo", "cfo"),
        ("LEV", "lev", "lev"),
        ("LOSS", "loss", "loss"),
        ("MB", "mb", "mb"),
        ("LIT", "lit", "lit"),
        ("ALTMAN", "altman", "altman"),
        ("TENURE", "tenure_ln", "tenure_ln"),
        ("ABS_ACCR_LAG", "ta_lag_abs", "ta_lag_abs"),
        ("BIG4", "big4", "big4"),
        ("SEC_TIER", "sec_tier", "sec_tier"),
        ("National Specialist", "nat_spec_d1", "nat_spec_d2"),
        ("City Specialist", "city_spec_d1", "city_spec_d2"),
        ("Both National and City Specialist", "both_d1", "both_d2"),
        ("National Specialist Only", "nat_only_d1", "nat_only_d2"),
        ("City Specialist Only", "city_only_d1", "city_only_d2"),
    ]
    rows: List[Tuple[str, List[str]]] = []
    for label, t1, t2 in spec:
        rows.append((label, row_values(df, t1, t2)))

    f_vals: List[str] = []
    adj_vals: List[str] = []
    for d, m in MODEL_ORDER:
        fv, fp, adj = model_stat_t5(df, d, m)
        f_vals.extend([fmt_est(fv), fmt_p(fp)])
        adj_vals.extend([fmt_est(adj), ""])
    rows.append(("F-value", f_vals))
    rows.append(("Adj. R2", adj_vals))
    return rows


def build_table8_rows(df: pd.DataFrame) -> List[Tuple[str, List[str]]]:
    spec = [
        ("Intercept", "Intercept", "Intercept"),
        ("SIZE", "size", "size"),
        ("sigma(EARN)", "sigma_earn", "sigma_earn"),
        ("LEV", "lev", "lev"),
        ("LOSS", "loss", "loss"),
        ("ROA", "roa", "roa"),
        ("MB", "mb", "mb"),
        ("LIT", "lit", "lit"),
        ("ALTMAN", "altman", "altman"),
        ("TENURE", "tenure_ln", "tenure_ln"),
        ("ACCR", "accr", "accr"),
        ("BIG4", "big4", "big4"),
        ("SEC_TIER", "sec_tier", "sec_tier"),
        ("National Specialist", "nat_spec_d1", "nat_spec_d2"),
        ("City Specialist", "city_spec_d1", "city_spec_d2"),
        ("Both National and City Specialist", "both_d1", "both_d2"),
        ("National Specialist Only", "nat_only_d1", "nat_only_d2"),
        ("City Specialist Only", "city_only_d1", "city_only_d2"),
    ]
    rows: List[Tuple[str, List[str]]] = []
    for label, t1, t2 in spec:
        rows.append((label, row_values(df, t1, t2)))

    lr_vals: List[str] = []
    pr2_vals: List[str] = []
    for d, m in MODEL_ORDER:
        lr, lrp, pr2 = model_stat_t8(df, d, m)
        lr_vals.extend([fmt_est(lr), fmt_p(lrp)])
        pr2_vals.extend([fmt_est(pr2), ""])
    rows.append(("Likelihood ratio", lr_vals))
    rows.append(("Pseudo-R2", pr2_vals))
    return rows


def main() -> None:
    args = parse_args()
    t5 = pd.read_csv(args.table5_fullcoef)
    t8 = pd.read_csv(args.table8_fullcoef)

    doc = Document()
    apply_font(doc)

    add_table_block(
        doc,
        "TABLE 5 (Replication)",
        "Dependent variable is the absolute value of abnormal accruals.",
        build_table5_rows(t5),
        "Coefficient p-values are two-tailed and based on cluster-robust inference by company_fkey. "
        "All continuous variables in equations (2)-(6) are winsorized at the 1st and 99th percentiles.",
    )
    doc.add_page_break()
    add_table_block(
        doc,
        "TABLE 8 (Replication)",
        "Dependent variable is the probability of issuing a going-concern opinion (GC).",
        build_table8_rows(t8),
        "Coefficient p-values are two-tailed and based on cluster-robust inference by company_fkey. "
        "All continuous variables in equations (2)-(6) are winsorized at the 1st and 99th percentiles.",
    )

    out = Path(args.output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
