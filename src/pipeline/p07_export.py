#!/usr/bin/env python3
"""
p07_export.py — 第七步：输出论文格式表格
==========================================
将 Table 5 和 Table 8 的回归结果格式化为：
  - Markdown（便于预览）
  - LaTeX（直接嵌入论文）
  - 完整诊断 JSON

格式严格对齐论文原表：
  Table 5: 6 列 (Def1 M1/M2/M3 + Def2 M1/M2/M3), 每列 Estimate + p-value
  Table 8: 同上

输入:
  outputs/table5_replication_fullcoef.csv
  outputs/table8_replication_fullcoef.csv
  data/processed/pipeline/s04_attrition.json
  data/processed/pipeline/s05_counts.json
  data/processed/pipeline/s06_counts.json
输出:
  outputs/table5_replication_paperstyle.md
  outputs/table5_replication_paperstyle.tex
  outputs/table8_replication_paperstyle.md
  outputs/table8_replication_paperstyle.tex
  outputs/replication_diagnostics.json
"""
from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches, Cm, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from pipeline.config import PipelineConfig, ensure_dirs
from pipeline.utils import fmt_est, fmt_p, latex_escape


# ═══════════════════════════════════════════════════════════════════
# 表格构建
# ═══════════════════════════════════════════════════════════════════

def _load_results(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def _get_est_p(full_df: pd.DataFrame,
               dnum: int, model: str, term: str) -> Tuple[str, str]:
    """从 fullcoef 表中提取某 definition-model-term 的 estimate 和 p-value。"""
    mask = (
        (full_df["definition"] == dnum)
        & (full_df["model"] == model)
        & (full_df["term"] == term)
    )
    sub = full_df.loc[mask]
    if sub.empty:
        return "", ""
    row = sub.iloc[0]
    return fmt_est(row.get("coef")), fmt_p(row.get("pvalue"))


def _get_fit(full_df: pd.DataFrame,
             dnum: int, model: str, label: str) -> Tuple[str, str]:
    """提取拟合优度指标。"""
    mask = (full_df["definition"] == dnum) & (full_df["model"] == model)
    sub = full_df.loc[mask]
    if sub.empty:
        return "", ""
    row = sub.iloc[0]
    if label == "F-value":
        return (fmt_est(row.get("f_value", np.nan)),
                fmt_p(row.get("f_pvalue", np.nan)))
    if label == "Adj. R2":
        return fmt_est(row.get("adj_r2", np.nan)), ""
    if label == "Likelihood ratio":
        return (fmt_est(row.get("lr_stat", np.nan)),
                fmt_p(row.get("lr_pvalue", np.nan)))
    if label == "Pseudo-R2":
        return fmt_est(row.get("pseudo_r2", np.nan)), ""
    return "", ""


def build_paper_df(
    row_specs: List[Tuple[str, object]],
    full_df: pd.DataFrame,
    fit_labels: List[str],
) -> pd.DataFrame:
    """构建与论文格式一致的宽表。"""
    model_order = [
        (1, "model1"), (1, "model2"), (1, "model3"),
        (2, "model1"), (2, "model2"), (2, "model3"),
    ]
    cols = ["Variable"]
    for d, m in model_order:
        mnum = m[-1]
        cols += [f"D{d}M{mnum}_Est", f"D{d}M{mnum}_p"]

    rows = []
    for label, term_spec in row_specs:
        out = {"Variable": label}
        for d, m in model_order:
            # term 可以是 str 或 {definition: str}
            if isinstance(term_spec, dict):
                term = term_spec.get(d, "")
            else:
                term = str(term_spec)
            est_col = f"D{d}M{m[-1]}_Est"
            p_col = f"D{d}M{m[-1]}_p"
            if term:
                est, pv = _get_est_p(full_df, d, m, term)
            else:
                est, pv = "", ""
            out[est_col] = est
            out[p_col] = pv
        rows.append(out)

    for label in fit_labels:
        out = {"Variable": label}
        for d, m in model_order:
            est_col = f"D{d}M{m[-1]}_Est"
            p_col = f"D{d}M{m[-1]}_p"
            est, pv = _get_fit(full_df, d, m, label)
            out[est_col] = est
            out[p_col] = pv
        rows.append(out)

    return pd.DataFrame(rows, columns=cols)


def write_md(path: Path, title: str, subtitle: str,
             df: pd.DataFrame) -> None:
    lines = [f"# {title}", "", subtitle, "", df.to_markdown(index=False)]
    path.write_text("\n".join(lines))


def write_tex(path: Path, title: str, subtitle: str,
              df: pd.DataFrame, note: str) -> None:
    model_headers = [
        (1, "Model 1"), (1, "Model 2"), (1, "Model 3"),
        (2, "Model 1"), (2, "Model 2"), (2, "Model 3"),
    ]
    lines = []
    lines.append(r"\begin{table}[!htbp]\centering")
    lines.append(r"\scriptsize")
    lines.append(rf"\caption{{{latex_escape(title)}}}")
    lines.append(rf"\textit{{{latex_escape(subtitle)}}}")
    lines.append(r"\begin{tabular}{l" + "rr" * 6 + "}")
    lines.append(r"\toprule")
    lines.append(
        r"& \multicolumn{6}{c}{Definition 1}"
        r" & \multicolumn{6}{c}{Definition 2}\\"
    )
    lines.append(r"\cmidrule(lr){2-7}\cmidrule(lr){8-13}")

    head2 = [" "]
    for _, mname in model_headers:
        head2.extend([mname, ""])
    lines.append(" & ".join(head2) + r"\\")
    lines.append(
        r"\cmidrule(lr){2-3}\cmidrule(lr){4-5}\cmidrule(lr){6-7}"
        r"\cmidrule(lr){8-9}\cmidrule(lr){10-11}\cmidrule(lr){12-13}"
    )
    head3 = ["Variable"] + ["Estimate", "p-value"] * 6
    lines.append(" & ".join(head3) + r"\\")
    lines.append(r"\midrule")

    for _, r in df.iterrows():
        vals = [latex_escape(str(r["Variable"]))]
        for d, m in [(1, "model1"), (1, "model2"), (1, "model3"),
                     (2, "model1"), (2, "model2"), (2, "model3")]:
            est = str(r[f"D{d}M{m[-1]}_Est"]).replace("<", "$<$")
            pv = str(r[f"D{d}M{m[-1]}_p"]).replace("<", "$<$")
            if est == "nan":
                est = ""
            if pv == "nan":
                pv = ""
            vals.extend([est, pv])
        lines.append(" & ".join(vals) + r"\\")

    lines.append(r"\bottomrule")
    lines.append(r"\end{tabular}")
    lines.append(
        rf"\vspace{{0.3em}}\par\footnotesize{{{latex_escape(note)}}}"
    )
    lines.append(r"\end{table}")
    path.write_text("\n".join(lines))


# ═══════════════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════════════

def run(cfg: PipelineConfig) -> None:
    ensure_dirs(cfg)
    out_dir = Path(cfg.output_dir)
    pipe = Path(cfg.pipeline_dir)

    # ── Table 5 ─────────────────────────────────────────────────
    print("[p07] 格式化 Table 5 ...")
    t5_full = _load_results(out_dir / "table5_replication_fullcoef.csv")

    t5_rows = [
        ("Intercept", "const"),
        ("SIZE", "size"),
        ("σ(CFO)", "sigma_cfo"),
        ("CFO", "cfo"),
        ("LEV", "lev"),
        ("LOSS", "loss"),
        ("MB", "mb"),
        ("LIT", "lit"),
        ("ALTMAN", "altman"),
        ("TENURE", "tenure_ln"),
        ("|ACCR_{t-1}|", "ta_lag_abs"),
        ("BIG4", "big4"),
        ("SEC_TIER", "sec_tier"),
        ("National Specialist",
         {1: "nat_spec_d1", 2: "nat_spec_d2"}),
        ("City Specialist",
         {1: "city_spec_d1", 2: "city_spec_d2"}),
        ("Both National and City Specialist",
         {1: "both_d1", 2: "both_d2"}),
        ("National Specialist Only",
         {1: "nat_only_d1", 2: "nat_only_d2"}),
        ("City Specialist Only",
         {1: "city_only_d1", 2: "city_only_d2"}),
    ]
    t5_df = build_paper_df(t5_rows, t5_full, ["F-value", "Adj. R2"])

    write_md(
        out_dir / "table5_replication_paperstyle.md",
        "TABLE 5 (Replication)",
        "Dependent variable: |DACC| (absolute value of abnormal accruals)",
        t5_df,
    )
    write_tex(
        out_dir / "table5_replication_paperstyle.tex",
        "TABLE 5 (Replication)",
        "Dependent variable: |DACC| (absolute value of abnormal accruals)",
        t5_df,
        "p-values are two-tailed, based on cluster-robust SE by firm.",
    )

    # ── Table 8 ─────────────────────────────────────────────────
    print("[p07] 格式化 Table 8 ...")
    t8_full = _load_results(out_dir / "table8_replication_fullcoef.csv")

    t8_rows = [
        ("Intercept", "Intercept"),
        ("SIZE", "size"),
        ("σ(EARN)", "sigma_earn"),
        ("LEV", "lev"),
        ("LOSS", "loss"),
        ("ROA", "roa"),
        ("MB", "mb"),
        ("LIT", "lit"),
        ("ALTMAN", "altman"),
        ("TENURE", "tenure_ln"),
        ("ACCR", "accr"),
        ("BIG4", "big4"),
        ("SEC_TIER", "sec_tier"),
        ("National Specialist",
         {1: "nat_spec_d1", 2: "nat_spec_d2"}),
        ("City Specialist",
         {1: "city_spec_d1", 2: "city_spec_d2"}),
        ("Both National and City Specialist",
         {1: "both_d1", 2: "both_d2"}),
        ("National Specialist Only",
         {1: "nat_only_d1", 2: "nat_only_d2"}),
        ("City Specialist Only",
         {1: "city_only_d1", 2: "city_only_d2"}),
    ]
    t8_df = build_paper_df(
        t8_rows, t8_full, ["Likelihood ratio", "Pseudo-R2"]
    )

    write_md(
        out_dir / "table8_replication_paperstyle.md",
        "TABLE 8 (Replication)",
        "Dependent variable: GC (going-concern opinion, logit)",
        t8_df,
    )
    write_tex(
        out_dir / "table8_replication_paperstyle.tex",
        "TABLE 8 (Replication)",
        "Dependent variable: GC (going-concern opinion, logit)",
        t8_df,
        "p-values are two-tailed, based on cluster-robust SE by firm."
        " Industry and year fixed effects included.",
    )

    # ── Word 输出（经济学表格样式） ───────────────────────────
    if HAS_DOCX:
        print("[p07] 格式化 Word 文档 ...")
        _write_word(
            out_dir / "table5_table8_replication_paperstyle.docx",
            t5_full, t5_rows,
            t8_full, t8_rows,
        )
    else:
        print("[p07] python-docx 未安装，跳过 Word 输出")

    # ── 综合诊断 ────────────────────────────────────────────────
    diag = {}
    for f in ["s04_attrition.json", "s05_counts.json", "s06_counts.json"]:
        p = pipe / f
        if p.exists():
            diag[f.replace(".json", "")] = json.loads(p.read_text())

    diag["notes"] = [
        "Office location: Audit Opinions AUDITOR_CITY/AUDITOR_STATE only.",
        "SIZE = ln(CSHO × PRCC_F).",
        "Eq(1) estimated on full Compustat (not limited to AA sample).",
        "Eq(1) vars truncated at 1%/99% before cross-sectional estimation.",
        "Eq(2)-(6) continuous vars winsorized at 1%/99% on analysis sample.",
        "Table 5: no industry/year FE (absorbed by cross-sectional Eq(1)).",
        "Table 8: two-digit SIC FE + fiscal year FE.",
        "Cluster-robust SE by company_fkey (Rogers 1993).",
    ]
    (out_dir / "replication_diagnostics.json").write_text(
        json.dumps(diag, indent=2, ensure_ascii=False)
    )

    print("[p07] 完成 → outputs/ 目录:")
    for f in sorted(out_dir.glob("table*_replication_paperstyle.*")):
        print(f"       {f.name}")
    print(f"       replication_diagnostics.json")


# ═══════════════════════════════════════════════════════════════════
# Word 经济学期刊样式表格
# ═══════════════════════════════════════════════════════════════════

def _stars(pvalue) -> str:
    """星号标注: *10%, **5%, ***1%."""
    if pd.isna(pvalue):
        return ""
    p = float(pvalue)
    if p < 0.01:
        return "***"
    if p < 0.05:
        return "**"
    if p < 0.10:
        return "*"
    return ""


def _get_coef_p(full_df: pd.DataFrame,
                dnum: int, model: str, term: str):
    mask = (
        (full_df["definition"] == dnum)
        & (full_df["model"] == model)
        & (full_df["term"] == term)
    )
    sub = full_df.loc[mask]
    if sub.empty:
        return np.nan, np.nan
    row = sub.iloc[0]
    return row.get("coef", np.nan), row.get("pvalue", np.nan)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框。kwargs: top/bottom/start/end = {sz, val, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_all_borders(table):
    """去掉表格全部边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    # 移除旧 borders
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(borders)


def _add_hline(table, row_idx, edge="bottom", sz="12"):
    """给某行所有单元格加水平线。"""
    for cell in table.rows[row_idx].cells:
        _set_cell_border(cell, **{edge: {"sz": sz, "val": "single", "color": "000000"}})


def _set_cell_font(cell, size=10, bold=False, italic=False, align="center"):
    for p in cell.paragraphs:
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _write_econ_table(
    doc: Document,
    title: str,
    subtitle: str,
    full_df: pd.DataFrame,
    row_specs,
    fit_rows: List[Tuple[str, str]],
    note: str,
):
    """写入一张经济学期刊样式表格。

    - 无竖线，仅顶线 / 表头下线 / 底线
    - 系数带星号，p-value 在下方括号内
    - Times New Roman 10pt
    """
    MODEL_ORDER = [
        (1, "model1"), (1, "model2"), (1, "model3"),
        (2, "model1"), (2, "model2"), (2, "model3"),
    ]

    # -- 标题 --
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(12)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(subtitle)
    r_sub.italic = True
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(10)

    # -- 计算行数 --
    # 每个变量占2行（系数+星号, p-value括号）
    # 表头: 3行 (def-group, model-group, Variable/Est列名)
    # fit_rows: 每个1行
    # N 行: 1行
    n_var_rows = len(row_specs) * 2
    n_fit = len(fit_rows)
    n_header = 3
    total_rows = n_header + n_var_rows + 1 + n_fit  # +1 for N row
    n_cols = 7  # Variable + 6 models

    table = doc.add_table(rows=total_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 去掉所有边框
    _remove_all_borders(table)

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        for j in range(1, 7):
            row.cells[j].width = Cm(2.0)

    # -- 表头 row 0: Definition 分组 --
    table.cell(0, 0).text = ""
    c1 = table.cell(0, 1)
    c1.merge(table.cell(0, 3)).text = "Definition 1"
    c2 = table.cell(0, 4)
    c2.merge(table.cell(0, 6)).text = "Definition 2"
    for c in range(7):
        _set_cell_font(table.cell(0, c), size=10, bold=True)

    # -- 表头 row 1: Model 分组 --
    table.cell(1, 0).text = ""
    for j, label in zip(range(1, 7), ["Model 1", "Model 2", "Model 3"] * 2):
        table.cell(1, j).text = label
        _set_cell_font(table.cell(1, j), size=9, bold=False, italic=True)
    _set_cell_font(table.cell(1, 0), size=10)

    # -- 表头 row 2: Variable 列名 --
    table.cell(2, 0).text = "Variable"
    _set_cell_font(table.cell(2, 0), size=10, bold=False, italic=True, align="left")
    for j in range(1, 7):
        table.cell(2, j).text = ""

    # 顶线（row 0 顶部）和表头下线（row 2 底部）
    _add_hline(table, 0, edge="top", sz="12")
    _add_hline(table, 2, edge="bottom", sz="6")

    # -- 数据行 --
    r_idx = n_header
    for label, term_spec in row_specs:
        # Row 1: 系数 + 星号
        table.cell(r_idx, 0).text = label
        _set_cell_font(table.cell(r_idx, 0), size=10, align="left")

        for j, (dnum, model) in enumerate(MODEL_ORDER):
            if isinstance(term_spec, dict):
                term = term_spec.get(dnum, "")
            else:
                term = str(term_spec)

            coef, pv = _get_coef_p(full_df, dnum, model, term) if term else (np.nan, np.nan)

            if pd.notna(coef):
                star = _stars(pv)
                table.cell(r_idx, j + 1).text = f"{float(coef):.4f}{star}"
            else:
                table.cell(r_idx, j + 1).text = ""
            _set_cell_font(table.cell(r_idx, j + 1), size=10)

        # Row 2: (p-value) 括号
        r_idx += 1
        table.cell(r_idx, 0).text = ""
        for j, (dnum, model) in enumerate(MODEL_ORDER):
            if isinstance(term_spec, dict):
                term = term_spec.get(dnum, "")
            else:
                term = str(term_spec)

            _, pv = _get_coef_p(full_df, dnum, model, term) if term else (np.nan, np.nan)

            if pd.notna(pv):
                pv_f = float(pv)
                pstr = f"<0.001" if pv_f < 0.001 else f"{pv_f:.3f}"
                table.cell(r_idx, j + 1).text = f"({pstr})"
            else:
                table.cell(r_idx, j + 1).text = ""
            _set_cell_font(table.cell(r_idx, j + 1), size=9, italic=False)

        r_idx += 1

    # -- N 行 --
    table.cell(r_idx, 0).text = "N"
    _set_cell_font(table.cell(r_idx, 0), size=10, align="left")
    for j, (dnum, model) in enumerate(MODEL_ORDER):
        mask = (full_df["definition"] == dnum) & (full_df["model"] == model)
        sub = full_df.loc[mask]
        if not sub.empty:
            n_obs = int(sub.iloc[0].get("n", 0))
            table.cell(r_idx, j + 1).text = f"{n_obs:,}"
        _set_cell_font(table.cell(r_idx, j + 1), size=10)
    r_idx += 1

    # -- Fit 行 (F-value / Adj R2 / LR / Pseudo R2) --
    for fit_label, fit_key in fit_rows:
        table.cell(r_idx, 0).text = fit_label
        _set_cell_font(table.cell(r_idx, 0), size=10, align="left")
        for j, (dnum, model) in enumerate(MODEL_ORDER):
            mask = (full_df["definition"] == dnum) & (full_df["model"] == model)
            sub = full_df.loc[mask]
            if not sub.empty:
                val = sub.iloc[0].get(fit_key, np.nan)
                if pd.notna(val):
                    table.cell(r_idx, j + 1).text = f"{float(val):.4f}"
            _set_cell_font(table.cell(r_idx, j + 1), size=10)
        r_idx += 1

    # 底线
    _add_hline(table, total_rows - 1, edge="bottom", sz="12")

    # -- 注释 --
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(4)
    r_note = p_note.add_run(note)
    r_note.font.name = "Times New Roman"
    r_note.font.size = Pt(9)
    r_note.italic = True


def _write_word(out_path: Path, t5_full, t5_rows, t8_full, t8_rows):
    """生成经济学期刊样式 Word 文档，包含 Table 5 和 Table 8。"""
    doc = Document()

    # 全局字体
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Table 5
    _write_econ_table(
        doc,
        title="TABLE 5",
        subtitle="OLS Regression of Absolute Abnormal Accruals on Auditor Industry Specialization",
        full_df=t5_full,
        row_specs=t5_rows,
        fit_rows=[("F-value", "f_value"), ("Adj. R²", "adj_r2")],
        note=(
            "Notes: ***, **, and * denote statistical significance at the 1%, 5%, and 10% levels "
            "(two-tailed), respectively. p-values in parentheses are based on standard errors "
            "clustered by firm (Rogers, 1993). All continuous variables are winsorized at the "
            "1st and 99th percentiles."
        ),
    )

    doc.add_page_break()

    # Table 8
    _write_econ_table(
        doc,
        title="TABLE 8",
        subtitle="Logit Regression of Going-Concern Opinions on Auditor Industry Specialization",
        full_df=t8_full,
        row_specs=t8_rows,
        fit_rows=[("Likelihood ratio", "lr_stat"), ("Pseudo R²", "pseudo_r2")],
        note=(
            "Notes: ***, **, and * denote statistical significance at the 1%, 5%, and 10% levels "
            "(two-tailed), respectively. p-values in parentheses are based on standard errors "
            "clustered by firm (Rogers, 1993). Industry (two-digit SIC) and year fixed effects "
            "are included but not reported. All continuous variables are winsorized at the "
            "1st and 99th percentiles. Sample restricted to financially distressed firms (OCF < 0)."
        ),
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[p07] Word → {out_path}")


if __name__ == "__main__":
    cfg = PipelineConfig.from_cli()
    run(cfg)
