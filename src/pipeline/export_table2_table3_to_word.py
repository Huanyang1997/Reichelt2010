#!/usr/bin/env python3
"""Export comprehensive Table 2 / Table 3 replication comparisons to Word."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


PAPER_BENCHMARKS_TABLE2: Dict[str, Dict[str, float]] = {
    "National Specialist": {"d1_mean": 0.116, "d1_sd": 0.320, "d2_mean": 0.214, "d2_sd": 0.410},
    "City Specialist": {"d1_mean": 0.350, "d1_sd": 0.477, "d2_mean": 0.327, "d2_sd": 0.469},
    "Both National and City Specialist": {"d1_mean": 0.076, "d1_sd": 0.265, "d2_mean": 0.124, "d2_sd": 0.330},
    "National Specialist Only": {"d1_mean": 0.040, "d1_sd": 0.195, "d2_mean": 0.090, "d2_sd": 0.287},
    "City Specialist Only": {"d1_mean": 0.274, "d1_sd": 0.446, "d2_mean": 0.203, "d2_sd": 0.402},
}

TABLE2_ROWS: List[Tuple[str, str, str]] = [
    ("National Specialist", "nat_spec_d1", "nat_spec_d2"),
    ("City Specialist", "city_spec_d1", "city_spec_d2"),
    ("Both National and City Specialist", "both_d1", "both_d2"),
    ("National Specialist Only", "nat_only_d1", "nat_only_d2"),
    ("City Specialist Only", "city_only_d1", "city_only_d2"),
]

PANEL_A_PAPER: Dict[str, Tuple[float, float]] = {
    "abs_dacc": (0.104, 0.134),
    "dacc": (0.005, 0.164),
    "size": (5.470, 2.286),
    "sigma_cfo": (0.195, 0.652),
    "cfo": (-0.001, 0.298),
    "lev": (0.168, 0.230),
    "loss": (0.395, 0.489),
    "mb": (3.022, 5.653),
    "lit": (0.258, 0.438),
    "altman": (-0.069, 7.324),
    "tenure_ln": (1.803, 0.737),
    "ta_lag_abs": (0.126, 0.207),
    "big4": (0.694, 0.461),
    "sec_tier": (0.102, 0.303),
}

PANEL_C_PAPER: Dict[str, Tuple[float, float]] = {
    "gc": (0.299, 0.458),
    "size": (3.700, 1.939),
    "sigma_earn": (3.311, 18.372),
    "lev": (0.200, 0.476),
    "loss": (0.859, 0.348),
    "roa": (-0.384, 0.395),
    "mb": (3.258, 16.872),
    "lit": (0.288, 0.453),
    "altman": (-1.524, 2.956),
    "tenure_ln": (1.702, 0.738),
    "accr": (-0.773, 3.949),
    "big4": (0.449, 0.497),
    "sec_tier": (0.106, 0.308),
}

TABLE2_B_PANEL_ORDER: Dict[str, List[str]] = {
    "B1": [
        "Total National Industry Specialists",
        "Total Industries",
        "Total Industry-Auditor combinations",
        "PWC",
        "EY",
        "DT",
        "KPMG",
        "GRANT THORNTON",
        "BDO SEIDMAN",
        "ALL OTHER",
    ],
    "B2": [
        "PWC",
        "EY",
        "DT",
        "KPMG",
        "GRANT THORNTON",
        "BDO SEIDMAN",
        "ALL OTHER",
        "Total City Industry Specialists",
        "Total Cities",
        "Total Industries",
        "Total City-Industry combinations",
        "Total City-Industry-Auditor combinations",
    ],
    "B3": [
        "Total National Industry Specialists",
        "Total Industries",
        "Total Industry-Auditor combinations",
        "PWC",
        "EY",
        "DT",
        "KPMG",
        "GRANT THORNTON",
        "BDO SEIDMAN",
        "ALL OTHER",
    ],
    "B4": [
        "PWC",
        "EY",
        "DT",
        "KPMG",
        "GRANT THORNTON",
        "BDO SEIDMAN",
        "ALL OTHER",
        "Total City Industry Specialists",
        "Total Cities",
        "Total Industries",
        "Total City-Industry combinations",
        "Total City-Industry-Auditor combinations",
    ],
}


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Export comprehensive Table 2/3 replication comparison to Word."
    )
    p.add_argument("--t5", default="data/processed/pipeline/s05_table5_sample.csv")
    p.add_argument("--t8", default="data/processed/pipeline/s06_table8_sample.csv")
    p.add_argument("--s04", default="data/processed/pipeline/s04_working_sample.csv")
    p.add_argument("--table2b-compare", default="outputs/table2_panelB_compare_vs_paper.csv")
    p.add_argument("--table2b-rep", default="outputs/table2_panelB_replication_counts.csv")
    p.add_argument("--output-docx", default="outputs/table2_table3_full_replication_compare.docx")
    p.add_argument("--output-table2-csv", default="outputs/table2_panelA_pipeline_compare.csv")
    p.add_argument("--output-table2b-csv", default="outputs/table2_panelB_pipeline_compare.csv")
    p.add_argument("--output-table3-csv", default="outputs/table3_panelA_C_pipeline_compare.csv")
    return p.parse_args()


def fmt(x: float) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.3f}"


def fmt_int_or_blank(x) -> str:
    if pd.isna(x):
        return ""
    return str(int(round(float(x))))


def apply_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def center_cell(cell) -> None:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc: Document, text: str, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold


def add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
        center_cell(t.cell(0, j))
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
            center_cell(t.cell(i, j))


def calc_stats(s: pd.Series) -> Dict[str, float]:
    s = pd.to_numeric(s, errors="coerce")
    return {
        "n": int(s.notna().sum()),
        "mean": float(s.mean()),
        "sd": float(s.std(ddof=1)),
        "p25": float(s.quantile(0.25)),
        "median": float(s.quantile(0.50)),
        "p75": float(s.quantile(0.75)),
    }


def ordered_values(values: Iterable[str], preferred_order: List[str]) -> List[str]:
    values = [str(v) for v in values]
    seen = set()
    out = []
    for v in preferred_order:
        if v in values and v not in seen:
            out.append(v)
            seen.add(v)
    for v in sorted(values):
        if v not in seen:
            out.append(v)
            seen.add(v)
    return out


def build_table2_panel_a_stats(df: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, float]] = []
    n = int(len(df))
    for label, c1, c2 in TABLE2_ROWS:
        s1 = calc_stats(df.get(c1))
        s2 = calc_stats(df.get(c2))
        pb = PAPER_BENCHMARKS_TABLE2.get(label, {})
        rows.append(
            {
                "variable": label,
                "N": n,
                "rep_d1_mean": s1["mean"],
                "rep_d1_sd": s1["sd"],
                "rep_d1_p25": s1["p25"],
                "rep_d1_median": s1["median"],
                "rep_d1_p75": s1["p75"],
                "paper_d1_mean": pb.get("d1_mean"),
                "paper_d1_sd": pb.get("d1_sd"),
                "diff_d1_mean": s1["mean"] - pb.get("d1_mean", float("nan")),
                "diff_d1_sd": s1["sd"] - pb.get("d1_sd", float("nan")),
                "rep_d2_mean": s2["mean"],
                "rep_d2_sd": s2["sd"],
                "rep_d2_p25": s2["p25"],
                "rep_d2_median": s2["median"],
                "rep_d2_p75": s2["p75"],
                "paper_d2_mean": pb.get("d2_mean"),
                "paper_d2_sd": pb.get("d2_sd"),
                "diff_d2_mean": s2["mean"] - pb.get("d2_mean", float("nan")),
                "diff_d2_sd": s2["sd"] - pb.get("d2_sd", float("nan")),
            }
        )
    return pd.DataFrame(rows)


def build_table3_stats(df: pd.DataFrame, paper: Dict[str, Tuple[float, float]], panel: str) -> pd.DataFrame:
    rows: List[Dict[str, float]] = []
    for var, (paper_mean, paper_sd) in paper.items():
        if var in df.columns:
            st = calc_stats(df[var])
        else:
            st = {
                "n": 0,
                "mean": float("nan"),
                "sd": float("nan"),
                "p25": float("nan"),
                "median": float("nan"),
                "p75": float("nan"),
            }
        rows.append(
            {
                "panel": panel,
                "variable": var,
                "N": st["n"],
                "rep_mean": st["mean"],
                "rep_sd": st["sd"],
                "rep_p25": st["p25"],
                "rep_median": st["median"],
                "rep_p75": st["p75"],
                "paper_mean": paper_mean,
                "paper_sd": paper_sd,
                "diff_mean": st["mean"] - paper_mean,
                "diff_sd": st["sd"] - paper_sd,
            }
        )
    return pd.DataFrame(rows)


def load_table2_panel_b_paper(compare_path: Path) -> pd.DataFrame:
    if compare_path.exists():
        d = pd.read_csv(compare_path)
        need = {"panel", "row", "year", "paper"}
        if need.issubset(d.columns):
            d["year"] = pd.to_numeric(d["year"], errors="coerce").astype("Int64")
            d["paper"] = pd.to_numeric(d["paper"], errors="coerce")
            return d[["panel", "row", "year", "paper"]].drop_duplicates(
                ["panel", "row", "year"], keep="first"
            )
    return pd.DataFrame(columns=["panel", "row", "year", "paper"])


def _auditor_bucket(v) -> str:
    try:
        k = int(float(v))
    except Exception:
        return "ALL OTHER"
    if k == 1:
        return "PWC"
    if k == 2:
        return "EY"
    if k == 3:
        return "DT"
    if k == 4:
        return "KPMG"
    if k == 5:
        return "GRANT THORNTON"
    if k == 6:
        return "BDO SEIDMAN"
    return "ALL OTHER"


def build_table2_panel_b_replication_from_s04(s04_path: Path) -> pd.DataFrame:
    if not s04_path.exists():
        return pd.DataFrame(columns=["panel", "row", "year", "value"])

    need_cols = [
        "fiscal_year",
        "sic2",
        "msa_code",
        "auditor_fkey",
        "nat_spec_d1",
        "nat_spec_d2",
        "city_spec_d1",
        "city_spec_d2",
    ]
    d = pd.read_csv(s04_path, usecols=[c for c in need_cols if c in pd.read_csv(s04_path, nrows=0).columns])
    for c in ["fiscal_year", "sic2", "auditor_fkey", "nat_spec_d1", "nat_spec_d2", "city_spec_d1", "city_spec_d2"]:
        if c in d.columns:
            d[c] = pd.to_numeric(d[c], errors="coerce")
    if "msa_code" in d.columns:
        d["msa_code"] = d["msa_code"].astype("string")
    d = d.dropna(subset=["fiscal_year", "sic2", "auditor_fkey"]).copy()
    d["fiscal_year"] = d["fiscal_year"].astype(int)
    d["sic2"] = d["sic2"].astype(int)
    d["auditor_fkey"] = d["auditor_fkey"].astype(int)
    d["auditor_bucket"] = d["auditor_fkey"].apply(_auditor_bucket)

    out_rows: List[Dict[str, object]] = []
    for year in sorted(d["fiscal_year"].unique()):
        dy = d[d["fiscal_year"] == year].copy()
        nat_total_ind = int(dy["sic2"].nunique())
        nat_total_comb = int(dy[["sic2", "auditor_fkey"]].drop_duplicates().shape[0])

        # Definition 1
        nat1 = dy[dy["nat_spec_d1"] == 1][["sic2", "auditor_fkey", "auditor_bucket"]].drop_duplicates()
        out_rows.extend(
            [
                {"panel": "B1", "row": "Total National Industry Specialists", "year": year, "value": int(len(nat1))},
                {"panel": "B1", "row": "Total Industries", "year": year, "value": nat_total_ind},
                {"panel": "B1", "row": "Total Industry-Auditor combinations", "year": year, "value": nat_total_comb},
            ]
        )
        for bucket in ["PWC", "EY", "DT", "KPMG", "GRANT THORNTON", "BDO SEIDMAN", "ALL OTHER"]:
            out_rows.append(
                {
                    "panel": "B1",
                    "row": bucket,
                    "year": year,
                    "value": int((nat1["auditor_bucket"] == bucket).sum()),
                }
            )

        # Definition 2
        nat2 = dy[dy["nat_spec_d2"] == 1][["sic2", "auditor_fkey", "auditor_bucket"]].drop_duplicates()
        out_rows.extend(
            [
                {"panel": "B3", "row": "Total National Industry Specialists", "year": year, "value": int(len(nat2))},
                {"panel": "B3", "row": "Total Industries", "year": year, "value": nat_total_ind},
                {"panel": "B3", "row": "Total Industry-Auditor combinations", "year": year, "value": nat_total_comb},
            ]
        )
        for bucket in ["PWC", "EY", "DT", "KPMG", "GRANT THORNTON", "BDO SEIDMAN", "ALL OTHER"]:
            out_rows.append(
                {
                    "panel": "B3",
                    "row": bucket,
                    "year": year,
                    "value": int((nat2["auditor_bucket"] == bucket).sum()),
                }
            )

        # City-level requires msa_code
        city_base = dy[dy["msa_code"].notna() & dy["msa_code"].astype(str).str.strip().ne("")].copy()
        city_total = int(city_base["msa_code"].nunique())
        city_ind = int(city_base["sic2"].nunique())
        city_ind_comb = int(city_base[["msa_code", "sic2"]].drop_duplicates().shape[0])
        city_ind_aud_comb = int(city_base[["msa_code", "sic2", "auditor_fkey"]].drop_duplicates().shape[0])

        city1 = city_base[city_base["city_spec_d1"] == 1][
            ["msa_code", "sic2", "auditor_fkey", "auditor_bucket"]
        ].drop_duplicates()
        out_rows.extend(
            [
                {"panel": "B2", "row": "Total City Industry Specialists", "year": year, "value": int(len(city1))},
                {"panel": "B2", "row": "Total Cities", "year": year, "value": city_total},
                {"panel": "B2", "row": "Total Industries", "year": year, "value": city_ind},
                {"panel": "B2", "row": "Total City-Industry combinations", "year": year, "value": city_ind_comb},
                {"panel": "B2", "row": "Total City-Industry-Auditor combinations", "year": year, "value": city_ind_aud_comb},
            ]
        )
        for bucket in ["PWC", "EY", "DT", "KPMG", "GRANT THORNTON", "BDO SEIDMAN", "ALL OTHER"]:
            out_rows.append(
                {
                    "panel": "B2",
                    "row": bucket,
                    "year": year,
                    "value": int((city1["auditor_bucket"] == bucket).sum()),
                }
            )

        city2 = city_base[city_base["city_spec_d2"] == 1][
            ["msa_code", "sic2", "auditor_fkey", "auditor_bucket"]
        ].drop_duplicates()
        out_rows.extend(
            [
                {"panel": "B4", "row": "Total City Industry Specialists", "year": year, "value": int(len(city2))},
                {"panel": "B4", "row": "Total Cities", "year": year, "value": city_total},
                {"panel": "B4", "row": "Total Industries", "year": year, "value": city_ind},
                {"panel": "B4", "row": "Total City-Industry combinations", "year": year, "value": city_ind_comb},
                {"panel": "B4", "row": "Total City-Industry-Auditor combinations", "year": year, "value": city_ind_aud_comb},
            ]
        )
        for bucket in ["PWC", "EY", "DT", "KPMG", "GRANT THORNTON", "BDO SEIDMAN", "ALL OTHER"]:
            out_rows.append(
                {
                    "panel": "B4",
                    "row": bucket,
                    "year": year,
                    "value": int((city2["auditor_bucket"] == bucket).sum()),
                }
            )

    out = pd.DataFrame(out_rows)
    if out.empty:
        return pd.DataFrame(columns=["panel", "row", "year", "value"])
    return out.drop_duplicates(["panel", "row", "year"], keep="first")


def build_table2_panel_b_compare(
    s04_path: Path, compare_path: Path, rep_fallback_path: Path
) -> pd.DataFrame:
    rep = build_table2_panel_b_replication_from_s04(s04_path)
    paper = load_table2_panel_b_paper(compare_path)

    if rep.empty and rep_fallback_path.exists():
        d = pd.read_csv(rep_fallback_path)
        if {"panel", "row", "year", "value"}.issubset(d.columns):
            d["year"] = pd.to_numeric(d["year"], errors="coerce").astype("Int64")
            d["value"] = pd.to_numeric(d["value"], errors="coerce")
            rep = d[["panel", "row", "year", "value"]].drop_duplicates(
                ["panel", "row", "year"], keep="first"
            )

    if rep.empty and paper.empty:
        return pd.DataFrame(columns=["panel", "row", "year", "paper", "value", "diff"])

    out = rep.merge(paper, on=["panel", "row", "year"], how="outer")
    out["diff"] = pd.to_numeric(out["value"], errors="coerce") - pd.to_numeric(out["paper"], errors="coerce")
    out = out[["panel", "row", "year", "paper", "value", "diff"]].copy()
    out["year"] = pd.to_numeric(out["year"], errors="coerce").astype("Int64")
    return out.sort_values(["panel", "row", "year"]).reset_index(drop=True)


def add_table2_panel_a(doc: Document, table2_df: pd.DataFrame) -> None:
    add_para(doc, "Panel A: Specialist indicator descriptive statistics", bold=True)

    headers = [
        "Variable",
        "N",
        "D1 Paper Mean",
        "D1 Rep Mean",
        "D1 Diff",
        "D1 Paper SD",
        "D1 Rep SD",
        "D1 Diff",
        "D2 Paper Mean",
        "D2 Rep Mean",
        "D2 Diff",
        "D2 Paper SD",
        "D2 Rep SD",
        "D2 Diff",
    ]
    rows: List[List[str]] = []
    for _, r in table2_df.iterrows():
        rows.append(
            [
                str(r["variable"]),
                str(int(r["N"])),
                fmt(r["paper_d1_mean"]),
                fmt(r["rep_d1_mean"]),
                fmt(r["diff_d1_mean"]),
                fmt(r["paper_d1_sd"]),
                fmt(r["rep_d1_sd"]),
                fmt(r["diff_d1_sd"]),
                fmt(r["paper_d2_mean"]),
                fmt(r["rep_d2_mean"]),
                fmt(r["diff_d2_mean"]),
                fmt(r["paper_d2_sd"]),
                fmt(r["rep_d2_sd"]),
                fmt(r["diff_d2_sd"]),
            ]
        )
    add_table(doc, headers, rows)

    add_para(doc, "Panel A replication quantiles (P25 / Median / P75):")
    q_headers = [
        "Variable",
        "D1 P25",
        "D1 Median",
        "D1 P75",
        "D2 P25",
        "D2 Median",
        "D2 P75",
    ]
    q_rows = []
    for _, r in table2_df.iterrows():
        q_rows.append(
            [
                str(r["variable"]),
                fmt(r["rep_d1_p25"]),
                fmt(r["rep_d1_median"]),
                fmt(r["rep_d1_p75"]),
                fmt(r["rep_d2_p25"]),
                fmt(r["rep_d2_median"]),
                fmt(r["rep_d2_p75"]),
            ]
        )
    add_table(doc, q_headers, q_rows)


def add_table2_panel_b(doc: Document, panel_b_df: pd.DataFrame) -> None:
    if panel_b_df.empty:
        add_para(doc, "Panel B data not available.", italic=True)
        return

    add_para(doc, "Panel B: Year-by-year specialist count comparison", bold=True)
    panel_order = ordered_values(panel_b_df["panel"].dropna().unique(), ["B1", "B2", "B3", "B4"])
    for panel in panel_order:
        sub = panel_b_df[panel_b_df["panel"] == panel].copy()
        if sub.empty:
            continue
        years = sorted([int(y) for y in sub["year"].dropna().unique()])
        rows_with_paper = sub.loc[sub["paper"].notna(), "row"].dropna().unique().tolist()
        row_source = rows_with_paper if len(rows_with_paper) > 0 else sub["row"].dropna().unique()
        rows_order = ordered_values(row_source, TABLE2_B_PANEL_ORDER.get(panel, []))

        add_para(doc, f"Panel {panel}")
        headers = ["Row"]
        for y in years:
            headers.extend([f"{y} Paper", f"{y} Rep", f"{y} Diff"])

        body: List[List[str]] = []
        for row_name in rows_order:
            row_vals = [row_name]
            for y in years:
                v = sub[(sub["row"] == row_name) & (sub["year"] == y)]
                if v.empty:
                    row_vals.extend(["", "", ""])
                else:
                    rec = v.iloc[0]
                    row_vals.extend(
                        [
                            fmt_int_or_blank(rec.get("paper")),
                            fmt_int_or_blank(rec.get("value")),
                            fmt_int_or_blank(rec.get("diff")),
                        ]
                    )
            body.append(row_vals)
        add_table(doc, headers, body)


def add_table3_panel(doc: Document, panel_df: pd.DataFrame, panel_name: str) -> None:
    add_para(doc, f"Panel {panel_name}", bold=True)
    headers = [
        "Variable",
        "N",
        "Paper Mean",
        "Rep Mean",
        "Diff Mean",
        "Paper SD",
        "Rep SD",
        "Diff SD",
        "Rep P25",
        "Rep Median",
        "Rep P75",
    ]
    rows: List[List[str]] = []
    for _, r in panel_df.iterrows():
        rows.append(
            [
                str(r["variable"]),
                str(int(r["N"])),
                fmt(r["paper_mean"]),
                fmt(r["rep_mean"]),
                fmt(r["diff_mean"]),
                fmt(r["paper_sd"]),
                fmt(r["rep_sd"]),
                fmt(r["diff_sd"]),
                fmt(r["rep_p25"]),
                fmt(r["rep_median"]),
                fmt(r["rep_p75"]),
            ]
        )
    add_table(doc, headers, rows)


def main() -> None:
    args = parse_args()

    t5 = pd.read_csv(args.t5)
    t8 = pd.read_csv(args.t8)
    table2_panel_b_df = build_table2_panel_b_compare(
        s04_path=Path(args.s04),
        compare_path=Path(args.table2b_compare),
        rep_fallback_path=Path(args.table2b_rep),
    )

    table2_df = build_table2_panel_a_stats(t5)
    table3_a_df = build_table3_stats(t5, PANEL_A_PAPER, panel="A")
    table3_c_df = build_table3_stats(t8, PANEL_C_PAPER, panel="C")
    table3_df = pd.concat([table3_a_df, table3_c_df], ignore_index=True)

    out2 = Path(args.output_table2_csv)
    out2.parent.mkdir(parents=True, exist_ok=True)
    table2_df.to_csv(out2, index=False)

    out2b = Path(args.output_table2b_csv)
    out2b.parent.mkdir(parents=True, exist_ok=True)
    table2_panel_b_df.to_csv(out2b, index=False)

    out3 = Path(args.output_table3_csv)
    out3.parent.mkdir(parents=True, exist_ok=True)
    table3_df.to_csv(out3, index=False)

    doc = Document()
    apply_font(doc)

    add_heading(doc, "TABLE 2 (Full Replication Comparison)", level=1)
    add_para(
        doc,
        "Panel A and Panel B are reported to facilitate side-by-side comparison with the paper.",
        italic=True,
    )
    add_table2_panel_a(doc, table2_df)
    add_para(doc, "")
    add_table2_panel_b(doc, table2_panel_b_df)

    doc.add_page_break()
    add_heading(doc, "TABLE 3 (Replication vs Paper)", level=1)
    add_para(doc, "Descriptive statistics for regression samples.", italic=True)
    add_table3_panel(doc, table3_a_df, panel_name="A")
    add_para(doc, "")
    add_table3_panel(doc, table3_c_df, panel_name="C")

    out = Path(args.output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    print(out)
    print(out2)
    print(out2b)
    print(out3)


if __name__ == "__main__":
    main()
